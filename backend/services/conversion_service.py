"""
File Conversion Service
Integrates the production-grade file_converter pipeline.

PDF → DOCX pipeline (from reference implementation):
  Stage 1 - Page analysis  (text-based vs scanned via pdfplumber)
  Stage 2 - Text + table extraction  (pdfplumber, lines_strict tables only)
  Stage 3 - Image extraction  (pdfimages lossless → pdf2image crop fallback)
  Stage 4 - DOCX assembly  (python-docx, heading heuristics, inline bold/italic)

DOCX → PDF pipeline:
  1. LibreOffice headless  (best fidelity, all platforms)
  2. docx2pdf             (Windows/macOS with Word installed)
  3. Pandoc + WeasyPrint  (cross-platform fallback)
  4. ReportLab            (last resort, text only)

Install:
    pip install python-docx pdfplumber pypdf pdf2image Pillow pytesseract reportlab docx2pdf
    # Linux system tools:
    sudo apt-get install -y poppler-utils libreoffice tesseract-ocr
"""

import asyncio
import logging
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# ── PDF → DOCX ────────────────────────────────────────────────────────────────

def _convert_pdf_to_docx(pdf_bytes: bytes) -> bytes:
    """
    Full pipeline from reference file_converter.py.
    Writes to temp files (pdfimages/pdf2image need file paths), returns DOCX bytes.
    """
    import io
    import pdfplumber
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    warnings: list[str] = []

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir = Path(tmpdir)

        # Write PDF to temp file (pdfimages needs a real path)
        pdf_path = tmpdir / "input.pdf"
        pdf_path.write_bytes(pdf_bytes)

        # ── Stage 1: Analyse pages ─────────────────────────────────────────
        page_types: list[str] = []
        with pdfplumber.open(str(pdf_path)) as pdf:
            total_pages = len(pdf.pages)
            for page in pdf.pages:
                txt = page.extract_text() or ""
                page_types.append("text" if len(txt.strip()) > 30 else "scanned")

        logger.info("PDF: %d pages (%d text, %d scanned)",
                    total_pages,
                    page_types.count("text"),
                    page_types.count("scanned"))

        # ── Stage 2: Extract text, tables, image metadata ──────────────────
        pages_data: list[dict] = []
        with pdfplumber.open(str(pdf_path)) as pdf:
            for i, page in enumerate(pdf.pages):
                ptype = page_types[i]

                if ptype == "scanned":
                    text = _ocr_page(str(pdf_path), i, dpi=300, tmpdir=tmpdir)
                    if text:
                        warnings.append(f"Page {i+1}: OCR used (scanned page)")
                    else:
                        text = ""
                else:
                    text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""

                # Tables — lines_strict only, never text-based (avoids false positives)
                try:
                    strict = {
                        "vertical_strategy":      "lines_strict",
                        "horizontal_strategy":    "lines_strict",
                        "snap_tolerance":         3,
                        "join_tolerance":         3,
                        "edge_min_length":        3,
                        "min_words_vertical":     1,
                        "min_words_horizontal":   1,
                        "intersection_tolerance": 3,
                        "text_tolerance":         3,
                    }
                    table_bboxes = [t.bbox for t in page.find_tables(strict)]
                    tables       = page.extract_tables(strict) or []
                except Exception:
                    table_bboxes = []
                    tables       = []

                # Per-char data for heading/bold detection
                chars = page.chars or []
                avg_size = (
                    sum(c["size"] for c in chars) / len(chars)
                    if chars else 12.0
                )

                pages_data.append({
                    "index":        i,
                    "type":         ptype,
                    "text":         text,
                    "tables":       tables,
                    "table_bboxes": table_bboxes,
                    "images_meta":  page.images or [],
                    "chars":        chars,
                    "avg_font_size": avg_size,
                    "width":        float(page.width),
                    "height":       float(page.height),
                })

        # ── Stage 3: Extract images ────────────────────────────────────────
        img_dir = tmpdir / "images"
        img_dir.mkdir()
        extracted_images = _extract_images_pdfimages(str(pdf_path), str(img_dir))
        if not extracted_images:
            extracted_images = _extract_images_render_fallback(
                str(pdf_path), pages_data, dpi=300,
                tmpdir=tmpdir, warnings=warnings
            )

        # ── Stage 4: Assemble DOCX ─────────────────────────────────────────
        doc = Document()
        # Remove default empty paragraph
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)

        # Set page size to match first PDF page
        _set_page_size(doc, pages_data[0]["width"], pages_data[0]["height"])

        content_added = False

        for page_data in pages_data:
            i = page_data["index"]
            if i > 0:
                doc.add_page_break()

            table_bboxes = page_data["table_bboxes"]

            # Tables first (they appear at their Y positions; we interleave below)
            # Build ordered content list by Y position
            ordered: list[tuple[str, float, object]] = []

            # Text: use char-level grouping for font metadata
            chars_outside = [
                c for c in page_data["chars"]
                if not _in_bbox(c["x0"], c["top"], table_bboxes)
            ]
            for para in _chars_to_paragraphs(chars_outside):
                if para["text"].strip():
                    ordered.append(("text", para["y"], para))

            # Tables at their top Y position
            for j, tdata in enumerate(page_data["tables"]):
                y = table_bboxes[j][1] if j < len(table_bboxes) else 0
                ordered.append(("table", y, tdata))

            ordered.sort(key=lambda x: x[1])

            for item_type, _, item in ordered:
                if item_type == "text":
                    text = item["text"].strip()
                    if not text:
                        continue
                    text = re.sub(r"(\w)-\s+(\w)", r"\1\2", text)

                    font_size = item.get("font_size", 11.0)
                    is_bold   = item.get("bold", False)
                    is_italic = item.get("italic", False)

                    para = doc.add_paragraph()

                    # Heading heuristics (from reference _is_heading / _heading_level)
                    if _is_heading(text, font_size, page_data["avg_font_size"]):
                        level = _heading_level(text, font_size)
                        para.style = f"Heading {level}"
                        para.add_run(text)
                    else:
                        para.style     = "Normal"
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        run = para.add_run(text)
                        if is_bold:
                            run.bold = True
                        if is_italic:
                            run.italic = True
                        run.font.size = Pt(min(round(font_size), 12))

                    content_added = True

                elif item_type == "table":
                    rows = [r for r in item if r and any(c for c in (r or []))]
                    if not rows:
                        continue
                    col_count = max((len(r) for r in rows), default=0)
                    if not col_count:
                        continue
                    try:
                        tbl = doc.add_table(rows=len(rows), cols=col_count)
                        tbl.style = "Table Grid"
                        for r_idx, row in enumerate(rows):
                            for c_idx in range(col_count):
                                val  = row[c_idx] if c_idx < len(row) else ""
                                cell = tbl.cell(r_idx, c_idx)
                                cell.text = re.sub(
                                    r"(\w)-\s+(\w)", r"\1\2",
                                    str(val or "").strip()
                                )
                                if r_idx == 0:
                                    for pp in cell.paragraphs:
                                        for run in pp.runs:
                                            run.bold = True
                        doc.add_paragraph()   # spacer after table
                        content_added = True
                    except Exception as e:
                        logger.warning("Table render failed page %d: %s", i + 1, e)

            # Images for this page
            page_imgs = [img for img in extracted_images if img.get("page") == i]
            for img_info in page_imgs:
                try:
                    _add_image_paragraph(doc, img_info["path"])
                    content_added = True
                except Exception as ex:
                    warnings.append(f"Page {i+1}: image insert failed — {ex}")

        if not content_added:
            raise RuntimeError(
                "No content extracted. PDF may be fully scanned/image-only. "
                "Install tesseract + pdf2image for OCR: "
                "pip install pdf2image pytesseract && sudo apt install tesseract-ocr"
            )

        out_path = tmpdir / "output.docx"
        doc.save(str(out_path))
        result = out_path.read_bytes()

    if warnings:
        logger.info("PDF→DOCX warnings: %s", "; ".join(warnings))

    return result


# ── PDF→DOCX helpers ──────────────────────────────────────────────────────────

def _in_bbox(x: float, y: float, bboxes: list) -> bool:
    return any(
        tx0 - 1 <= x <= tx1 + 1 and ty0 - 1 <= y <= ty1 + 1
        for tx0, ty0, tx1, ty1 in bboxes
    )


def _is_heading(text: str, font_size: float, avg_size: float) -> bool:
    s = text.strip()
    if not s or len(s) > 120:
        return False
    # Significantly larger than page average
    if font_size >= avg_size * 1.3 and font_size >= 12:
        return True
    # Numbered section: "1 Introduction", "2.1 Methods"
    if re.match(r"^(\d+\.?){1,3}\s+[A-Z]", s):
        return True
    # ALL CAPS short line
    if s.isupper() and len(s) < 60:
        return True
    # Known section keywords
    if s.lower() in {
        "abstract", "introduction", "conclusion", "conclusions",
        "references", "bibliography", "acknowledgements", "acknowledgments",
        "methodology", "methods", "results", "discussion", "appendix",
        "related work", "background", "evaluation", "experiments",
        "overview", "summary", "approach", "implementation",
    }:
        return True
    return False


def _heading_level(text: str, font_size: float) -> int:
    if re.match(r"^\d+\.\d+\.\d+", text):
        return 3
    if re.match(r"^\d+\.\d+", text):
        return 2
    return 1


def _chars_to_paragraphs(chars: list) -> list[dict]:
    if not chars:
        return []

    chars = sorted(chars, key=lambda c: (round(c["top"], 1), c["x0"]))

    LINE_TOL = 2.5
    lines: list[list] = []
    current: list = []
    last_top = None

    for c in chars:
        if last_top is None or abs(c["top"] - last_top) <= LINE_TOL:
            current.append(c)
        else:
            if current:
                lines.append(current)
            current = [c]
        last_top = c["top"]
    if current:
        lines.append(current)

    paragraphs: list[dict] = []
    para_lines: list[list] = []
    last_bottom = None

    for line in lines:
        top    = min(c["top"]    for c in line)
        bottom = max(c["bottom"] for c in line)
        height = (bottom - top) or 10

        if last_bottom is not None and (top - last_bottom) > height * 1.4:
            if para_lines:
                paragraphs.append(_build_para(para_lines))
                para_lines = []

        para_lines.append(line)
        last_bottom = bottom

    if para_lines:
        paragraphs.append(_build_para(para_lines))

    return paragraphs


def _build_para(lines: list[list]) -> dict:
    all_chars = [c for ln in lines for c in ln]
    parts     = ["".join(c["text"] for c in ln) for ln in lines]
    text      = " ".join(parts)

    sizes  = [c.get("size", 11) for c in all_chars if c.get("size")]
    fonts  = [c.get("fontname", "") for c in all_chars]
    avg    = sum(sizes) / len(sizes) if sizes else 11
    bold   = any("Bold" in f or "bold" in f for f in fonts)
    italic = any("Italic" in f or "italic" in f or "Oblique" in f for f in fonts)

    return {
        "text":      text,
        "y":         min(c["top"] for c in all_chars),
        "font_size": round(avg, 1),
        "bold":      bold,
        "italic":    italic,
    }


def _set_page_size(doc, pdf_width_pts: float, pdf_height_pts: float):
    emu_per_pt      = 12700
    section         = doc.sections[0]
    section.page_width  = int(pdf_width_pts  * emu_per_pt)
    section.page_height = int(pdf_height_pts * emu_per_pt)


def _add_image_paragraph(doc, img_path: str):
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    try:
        from PIL import Image
        with Image.open(img_path) as img:
            w_px, _ = img.size
        w_in = min(5.0, w_px / 96)
    except Exception:
        w_in = 4.0
    para           = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(img_path, width=Inches(w_in))


def _extract_images_pdfimages(pdf_path: str, out_dir: str) -> list[dict]:
    if not shutil.which("pdfimages"):
        logger.info("pdfimages not found — skipping lossless image extraction")
        return []
    prefix = os.path.join(out_dir, "img")
    try:
        subprocess.run(["pdfimages", "-all", pdf_path, prefix],
                       capture_output=True, timeout=60)
    except Exception as e:
        logger.warning("pdfimages failed: %s", e)
        return []

    results = []
    for f in sorted(Path(out_dir).glob("img-*")):
        if f.suffix.lower() in (".png", ".jpg", ".jpeg", ".ppm", ".pbm", ".tif", ".tiff"):
            try:
                page_num = int(f.stem.split("-")[-2]) - 1
            except (ValueError, IndexError):
                page_num = 0
            results.append({"path": str(f), "page": page_num})

    logger.info("pdfimages: %d images extracted", len(results))
    return results


def _extract_images_render_fallback(
    pdf_path: str, pages_data: list, dpi: int,
    tmpdir: Path, warnings: list
) -> list[dict]:
    try:
        from pdf2image import convert_from_path
        from PIL import Image
    except ImportError:
        logger.info("pdf2image/Pillow not installed — image fallback skipped")
        return []

    renders_dir = tmpdir / "renders"
    renders_dir.mkdir(exist_ok=True)
    try:
        renders = convert_from_path(pdf_path, dpi=dpi)
    except Exception as e:
        warnings.append(f"Page rendering failed: {e}")
        return []

    results = []
    for i, (render, page_data) in enumerate(zip(renders, pages_data)):
        if not page_data["images_meta"]:
            continue
        pdf_w, pdf_h     = page_data["width"], page_data["height"]
        render_w, render_h = render.size
        sx, sy = render_w / pdf_w, render_h / pdf_h

        for j, meta in enumerate(page_data["images_meta"]):
            try:
                x0 = int(meta["x0"] * sx)
                y0 = int((pdf_h - meta["y1"]) * sy)
                x1 = int(meta["x1"] * sx)
                y1 = int((pdf_h - meta["y0"]) * sy)
                pad = 4
                x0, y0 = max(0, x0-pad), max(0, y0-pad)
                x1, y1 = min(render_w, x1+pad), min(render_h, y1+pad)
                if x1 <= x0 or y1 <= y0:
                    continue
                crop_path = str(renders_dir / f"page{i}_img{j}.png")
                render.crop((x0, y0, x1, y1)).save(crop_path, "PNG")
                results.append({"path": crop_path, "page": i})
            except Exception as ex:
                warnings.append(f"Page {i+1} image {j}: crop failed — {ex}")

    return results


def _ocr_page(pdf_path: str, page_index: int, dpi: int, tmpdir: Path) -> str:
    try:
        import pytesseract
        from pdf2image import convert_from_path
        renders = convert_from_path(
            pdf_path, dpi=dpi,
            first_page=page_index + 1,
            last_page=page_index + 1,
        )
        if renders:
            return pytesseract.image_to_string(renders[0], lang="eng")
    except ImportError:
        logger.info("pytesseract/pdf2image not installed — OCR skipped")
    except Exception as e:
        logger.warning("OCR failed page %d: %s", page_index + 1, e)
    return ""


# ── DOCX → PDF ────────────────────────────────────────────────────────────────

def _convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """
    Attempt DOCX→PDF in priority order:
    1. LibreOffice headless (best, all platforms)
    2. docx2pdf (Windows/macOS + Word)
    3. Pandoc + WeasyPrint
    4. ReportLab (last resort, text only)
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir    = Path(tmpdir)
        docx_path = tmpdir / "input.docx"
        pdf_path  = tmpdir / "output.pdf"
        docx_path.write_bytes(docx_bytes)

        # Method 1: LibreOffice
        lo = _find_libreoffice()
        if lo:
            try:
                r = subprocess.run(
                    [lo, "--headless", "--norestore", "--nofirststartwizard",
                     "--convert-to", "pdf", "--outdir", str(tmpdir), str(docx_path)],
                    capture_output=True, timeout=120,
                )
                # LibreOffice names output <stem>.pdf in outdir
                lo_out = tmpdir / "input.pdf"
                if lo_out.exists():
                    logger.info("DOCX→PDF via LibreOffice")
                    return lo_out.read_bytes()
            except Exception as e:
                logger.warning("LibreOffice failed: %s", e)

        # Method 2: docx2pdf
        try:
            from docx2pdf import convert as d2p
            d2p(str(docx_path), str(pdf_path))
            if pdf_path.exists():
                logger.info("DOCX→PDF via docx2pdf")
                return pdf_path.read_bytes()
        except ImportError:
            pass
        except Exception as e:
            logger.warning("docx2pdf failed: %s", e)

        # Method 3: Pandoc + WeasyPrint
        if shutil.which("pandoc"):
            try:
                from weasyprint import HTML
                html_path = tmpdir / "tmp.html"
                subprocess.run(
                    ["pandoc", str(docx_path), "-o", str(html_path),
                     "--embed-resources", "--standalone"],
                    check=True, capture_output=True, timeout=60,
                )
                HTML(filename=str(html_path)).write_pdf(str(pdf_path))
                if pdf_path.exists():
                    logger.info("DOCX→PDF via Pandoc + WeasyPrint")
                    return pdf_path.read_bytes()
            except ImportError:
                pass
            except Exception as e:
                logger.warning("Pandoc/WeasyPrint failed: %s", e)

        # Method 4: ReportLab (last resort)
        logger.warning("Using ReportLab fallback — text only, images replaced with placeholders")
        return _reportlab_fallback(str(docx_path), str(pdf_path))


def _find_libreoffice() -> Optional[str]:
    candidates = [
        "libreoffice", "soffice",
        "/usr/bin/libreoffice", "/usr/bin/soffice",
        "/usr/local/bin/libreoffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for c in candidates:
        if shutil.which(c) or Path(c).exists():
            return c
    return None


def _reportlab_fallback(docx_path: str, pdf_path: str) -> bytes:
    from docx import Document as DocxDocument
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer,
        Table, TableStyle, PageBreak,
    )

    doc_src = DocxDocument(docx_path)
    sec     = doc_src.sections[0]
    pw = (sec.page_width  / 914400) * 72 if sec.page_width  else 595
    ph = (sec.page_height / 914400) * 72 if sec.page_height else 842
    margin = 72
    content_width = pw - 2 * margin

    pdf_doc = SimpleDocTemplate(
        pdf_path, pagesize=(pw, ph),
        leftMargin=margin, rightMargin=margin,
        topMargin=margin, bottomMargin=margin,
    )

    styles = getSampleStyleSheet()
    body   = ParagraphStyle("body",   parent=styles["Normal"],  fontSize=11, leading=16, alignment=TA_JUSTIFY)
    h1     = ParagraphStyle("h1",     parent=styles["Heading1"], fontSize=18, leading=22, spaceAfter=12)
    h2     = ParagraphStyle("h2",     parent=styles["Heading2"], fontSize=14, leading=18, spaceAfter=8)
    h3     = ParagraphStyle("h3",     parent=styles["Heading3"], fontSize=12, leading=16, spaceAfter=6)
    cap    = ParagraphStyle("cap",    parent=styles["Normal"],  fontSize=9,  leading=12, alignment=TA_CENTER, textColor=colors.grey)
    hmap   = {1: h1, 2: h2, 3: h3}
    story  = []

    def esc(t):
        return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def rl_inline(para):
        parts = []
        for run in para.runs:
            t = esc(run.text)
            if run.bold:   t = f"<b>{t}</b>"
            if run.italic: t = f"<i>{t}</i>"
            parts.append(t)
        return "".join(parts) or esc(para.text)

    for para in doc_src.paragraphs:
        sname = para.style.name.lower()
        if sname.startswith("heading"):
            try: lvl = min(int(sname.split()[-1]), 3)
            except ValueError: lvl = 1
            txt = para.text.strip()
            if txt:
                story.append(Paragraph(esc(txt), hmap[lvl]))
                story.append(Spacer(1, 6))
            continue

        # Detect inline images
        has_img = any(
            (elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag) in ("drawing", "pict")
            for run in para.runs for elem in run._element
        )
        if has_img:
            story.append(Paragraph("[IMAGE]", cap))
            continue

        txt = para.text.strip()
        if not txt:
            story.append(Spacer(1, 6))
            continue
        story.append(Paragraph(rl_inline(para), body))

    for table in doc_src.tables:
        tdata = [[Paragraph(esc(cell.text), body) for cell in row.cells]
                 for row in table.rows]
        if tdata:
            ncols = len(tdata[0])
            cw    = content_width / ncols if ncols else content_width
            rl_t  = Table(tdata, colWidths=[cw] * ncols)
            rl_t.setStyle(TableStyle([
                ("GRID",       (0,0), (-1,-1), 0.5, colors.grey),
                ("BACKGROUND", (0,0), (-1, 0), colors.HexColor("#E8E8E8")),
                ("FONTSIZE",   (0,0), (-1,-1), 10),
                ("TOPPADDING", (0,0), (-1,-1), 4),
                ("BOTTOMPADDING", (0,0), (-1,-1), 4),
            ]))
            story.append(rl_t)
            story.append(Spacer(1, 10))

    pdf_doc.build(story)
    return Path(pdf_path).read_bytes()


# ── Public async wrappers ─────────────────────────────────────────────────────

async def pdf_to_docx(pdf_bytes: bytes, original_filename: str) -> tuple[bytes, str, str]:
    """Convert PDF → DOCX. Runs blocking I/O in thread executor."""
    stem = Path(original_filename).stem
    loop = asyncio.get_event_loop()
    try:
        docx_bytes = await loop.run_in_executor(None, _convert_pdf_to_docx, pdf_bytes)
    except Exception as e:
        logger.error("PDF→DOCX failed: %s", e)
        raise
    out_name = f"{stem}.docx"
    logger.info("PDF→DOCX: %s (%d → %d bytes)", original_filename, len(pdf_bytes), len(docx_bytes))
    return (
        docx_bytes,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        out_name,
    )


async def docx_to_pdf(docx_bytes: bytes, original_filename: str) -> tuple[bytes, str, str]:
    """Convert DOCX → PDF. Runs blocking I/O in thread executor."""
    stem = Path(original_filename).stem
    loop = asyncio.get_event_loop()
    try:
        pdf_out = await loop.run_in_executor(None, _convert_docx_to_pdf, docx_bytes)
    except Exception as e:
        logger.error("DOCX→PDF failed: %s", e)
        raise
    out_name = f"{stem}.pdf"
    logger.info("DOCX→PDF: %s (%d → %d bytes)", original_filename, len(docx_bytes), len(pdf_out))
    return pdf_out, "application/pdf", out_name
